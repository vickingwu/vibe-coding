"""Extract text from cal/PRD.docx and cal/UI.docx"""
from docx import Document
import os

for name in ['PRD.docx', 'UI.docx']:
    path = os.path.join('cal', name)
    print(f"\n{'='*60}")
    print(f"  {name}")
    print(f"{'='*60}\n")
    doc = Document(path)
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    # Also check tables
    for i, table in enumerate(doc.tables):
        print(f"\n[Table {i+1}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            print(' | '.join(cells))

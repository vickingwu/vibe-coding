# -*- coding: utf-8 -*-
"""
生成 LOF基金溢价率监控页面 设计与开发思路文档 (Word)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

output_dir = r'C:\Users\wuminqia\Documents\MQ'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'LOF溢价率监控_设计文档.docx')

doc = Document()

# 设置默认字体
st = doc.styles['Normal']
st.font.name = '微软雅黑'
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ============================================================
# 标题
# ============================================================
title = doc.add_heading('LOF基金溢价率监控页面 — 设计与开发思路', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('')

# 文档信息表
info_table = doc.add_table(rows=3, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('文档类型', '设计与开发思路说明'),
    ('日期', '2026-03-10'),
    ('状态', '待确认'),
]):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v

doc.add_paragraph('')

# ============================================================
# 1. 项目概述
# ============================================================
doc.add_heading('1. 项目概述', level=1)
doc.add_paragraph(
    '开发一个单页面 HTML 网页，用于实时监控 A 股市场中溢价率 > 0 的 LOF 基金。'
    '页面以表格形式展示基金数据，默认按 T-1 溢价率从高到低排序，'
    '支持一键刷新获取最新数据，以及点击列头切换排序方向。'
    'UI 风格参照 Amazon Cloudscape Design System。'
)

doc.add_paragraph('')

# ============================================================
# 2. 数据来源与获取方案
# ============================================================
doc.add_heading('2. 数据来源与获取方案', level=1)

doc.add_heading('2.1 数据来源', level=2)
doc.add_paragraph(
    '使用集思录（jisilu.cn）提供的 LOF 基金实时数据接口。'
    '该接口返回 JSON 格式数据，包含基金代码、名称、现价、净值、溢价率等完整字段。'
)

doc.add_heading('2.2 接口地址', level=2)
p = doc.add_paragraph()
r = p.add_run('https://www.jisilu.cn/data/lof/stock_lof_list/?___jsl=LST___t=...')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0, 115, 187)

doc.add_heading('2.3 跨域问题处理', level=2)
doc.add_paragraph(
    '由于浏览器的同源策略限制，前端无法直接请求集思录 API。方案如下：'
)
items = [
    ('方案 A（推荐）：CORS 代理', 
     '使用公共 CORS 代理服务（如 https://api.allorigins.win/raw?url=...）'
     '包装请求，绕过跨域限制。优点是纯前端实现，无需后端服务器。'),
    ('方案 B：本地代理服务器', 
     '用 Node.js 或 Python 起一个简单的本地代理服务，转发请求到集思录。'
     '适合长期稳定使用的场景。'),
    ('方案 C：手动数据模式', 
     '如果 API 不可用，页面内置一组示例数据，支持手动粘贴 JSON 更新。'
     '作为兜底方案。'),
]
for title_text, desc in items:
    p = doc.add_paragraph()
    r = p.add_run(title_text + '：')
    r.bold = True
    p.add_run(desc)

doc.add_paragraph('')
doc.add_paragraph(
    '实际开发中，三种方案都会实现。默认使用方案 A，'
    '如果代理不可用则自动降级到内置示例数据（方案 C），'
    '同时提供方案 B 的代码供本地部署。'
)

doc.add_paragraph('')

# ============================================================
# 3. 数据字段设计
# ============================================================
doc.add_heading('3. 数据字段设计', level=1)
doc.add_paragraph('表格展示以下字段：')

field_table = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
headers = ['字段', '来源/说明', '示例']
for i, h in enumerate(headers):
    field_table.rows[0].cells[i].text = h
    for p2 in field_table.rows[0].cells[i].paragraphs:
        for r2 in p2.runs:
            r2.bold = True

fields = [
    ('基金代码', 'fund_id，6位数字代码', '501029'),
    ('基金名称', 'fund_nm，基金全称', '华宝标普油气LOF'),
    ('相关标的', 'index_nm / 跟踪指数名称', '标普石油天然气上游股票指数'),
    ('现价', 'price，最新交易价格', '0.571'),
    ('T-1溢价率', 'estimate_value 或 discount_rt，T-1日溢价率', '3.25%'),
    ('市场分类', '根据跟踪标的自动分类：欧美/亚洲/商品/A股/其他', '商品'),
    ('申购状态', 'apply_status：开放申购/暂停申购/限额申购', '限500'),
    ('申购费率', 'apply_fee，申购费率百分比', '1.20%'),
]
for i, (f, s, e) in enumerate(fields):
    field_table.rows[i+1].cells[0].text = f
    field_table.rows[i+1].cells[1].text = s
    field_table.rows[i+1].cells[2].text = e

doc.add_paragraph('')

# ============================================================
# 4. 页面结构与布局
# ============================================================
doc.add_heading('4. 页面结构与布局', level=1)

doc.add_heading('4.1 整体布局', level=2)
doc.add_paragraph(
    '采用 Amazon Cloudscape 风格的单页面布局，无侧边栏，'
    '以内容为中心的简洁设计：'
)
layout_items = [
    '顶部导航栏：页面标题 + 数据更新时间 + 操作按钮',
    '统计卡片区：展示关键指标（溢价基金数量、最高溢价率、平均溢价率等）',
    '筛选/搜索区：市场分类筛选 + 申购状态筛选 + 关键词搜索',
    '数据表格区：核心内容，展示所有溢价率 > 0 的 LOF 基金',
    '页脚：数据来源说明 + 免责声明',
]
for item in layout_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 顶部区域', level=2)
doc.add_paragraph(
    '左侧显示页面标题"LOF基金溢价率监控"，'
    '右侧放置【查询最新数据】按钮（Primary 样式）。'
    '按钮点击后显示 loading 状态，数据加载完成后自动恢复。'
    '标题下方显示"数据更新时间：2026-03-10 14:30:00"。'
)

doc.add_heading('4.3 统计卡片', level=2)
doc.add_paragraph('顶部展示 4 个关键指标卡片：')
stats = [
    ('溢价基金数', '当前溢价率 > 0 的基金总数'),
    ('最高溢价率', '当前最高的 T-1 溢价率及对应基金名'),
    ('平均溢价率', '所有溢价基金的平均溢价率'),
    ('开放申购数', '溢价基金中申购状态为"开放申购"的数量'),
]
stats_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
stats_table.rows[0].cells[0].text = '指标'
stats_table.rows[0].cells[1].text = '说明'
for i, (k, v) in enumerate(stats):
    stats_table.rows[i+1].cells[0].text = k
    stats_table.rows[i+1].cells[1].text = v

doc.add_heading('4.4 筛选与搜索', level=2)
doc.add_paragraph(
    '表格上方提供筛选功能：\n'
    '• 市场分类下拉框：全部 / 欧美 / 亚洲 / 商品 / A股 / 其他\n'
    '• 申购状态下拉框：全部 / 开放申购 / 暂停申购 / 限额申购\n'
    '• 搜索框：支持按基金代码或名称模糊搜索\n'
    '筛选条件变化时实时过滤表格数据，无需刷新页面。'
)

doc.add_paragraph('')

# ============================================================
# 5. 交互设计
# ============================================================
doc.add_heading('5. 交互设计', level=1)

interactions = [
    ('【查询最新数据】按钮',
     '点击后：\n'
     '1. 按钮变为 loading 状态（显示旋转图标 + "加载中..."）\n'
     '2. 发起 API 请求获取最新数据\n'
     '3. 数据返回后过滤溢价率 > 0 的记录\n'
     '4. 按 T-1 溢价率降序排列\n'
     '5. 更新表格和统计卡片\n'
     '6. 更新"数据更新时间"\n'
     '7. 按钮恢复正常状态\n'
     '8. 如果请求失败，显示错误提示（Cloudscape Alert 样式）'),
    ('T-1溢价率列排序',
     '点击列头切换排序方向：\n'
     '• 默认降序（▼），最高溢价在最上面\n'
     '• 点击一次切换为升序（▲）\n'
     '• 再次点击切换回降序（▼）\n'
     '• 列头显示当前排序方向的箭头图标'),
    ('溢价率颜色编码',
     '根据溢价率数值显示不同颜色：\n'
     '• ≥ 5%：红色（高溢价警告）\n'
     '• 2% ~ 5%：橙色（中等溢价）\n'
     '• 0% ~ 2%：绿色（低溢价）'),
    ('申购状态标签',
     '使用 Badge 组件展示：\n'
     '• 开放申购：绿色 Badge\n'
     '• 暂停申购：红色 Badge\n'
     '• 限额申购（如限500）：橙色 Badge'),
    ('行悬停效果',
     '鼠标悬停在表格行上时，行背景变为浅灰色，增强可读性。'),
]

for title_text, desc in interactions:
    p = doc.add_paragraph()
    r = p.add_run(title_text)
    r.bold = True
    doc.add_paragraph(desc)
    doc.add_paragraph('')

# ============================================================
# 6. 市场分类逻辑
# ============================================================
doc.add_heading('6. 市场分类逻辑', level=1)
doc.add_paragraph(
    '根据基金跟踪的指数/标的名称，自动归类到以下市场：'
)

cat_table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
cat_table.rows[0].cells[0].text = '市场分类'
cat_table.rows[0].cells[1].text = '关键词匹配规则'
categories = [
    ('欧美', '标普、纳斯达克、道琼斯、美国、德国、DAX、法国、英国、MSCI'),
    ('亚洲', '日经、东证、恒生、印度、越南、韩国、东南亚、亚太'),
    ('商品', '黄金、白银、原油、油气、有色金属、豆粕、商品'),
    ('A股', '沪深300、中证500、创业板、科创板、上证50、中证1000'),
    ('其他', '不匹配以上任何关键词的基金'),
]
for i, (k, v) in enumerate(categories):
    cat_table.rows[i+1].cells[0].text = k
    cat_table.rows[i+1].cells[1].text = v

doc.add_paragraph('')

# ============================================================
# 7. UI 设计规范
# ============================================================
doc.add_heading('7. UI 设计规范（Amazon Cloudscape 风格）', level=1)

doc.add_heading('7.1 配色方案', level=2)
color_table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
color_table.rows[0].cells[0].text = '用途'
color_table.rows[0].cells[1].text = '色值'
color_table.rows[0].cells[2].text = '说明'
colors = [
    ('页面背景', '#f2f3f3', 'Cloudscape 标准背景色'),
    ('容器背景', '#ffffff', '白色卡片/表格背景'),
    ('主色调', '#0073bb', 'Cloudscape 蓝，用于按钮、链接'),
    ('文字主色', '#16191f', '正文文字'),
    ('文字辅助色', '#5f6b7a', '次要信息'),
    ('分割线', '#e9ebed', '表格边框、分割线'),
    ('强调色', '#ff9900', 'Amazon 橙，用于重要提示'),
]
for i, (u, c, d) in enumerate(colors):
    color_table.rows[i+1].cells[0].text = u
    color_table.rows[i+1].cells[1].text = c
    color_table.rows[i+1].cells[2].text = d

doc.add_heading('7.2 字体', level=2)
doc.add_paragraph(
    '字体栈：Amazon Ember, "Helvetica Neue", Roboto, Arial, sans-serif\n'
    '正文字号：14px\n'
    '表头字号：13px，大写字母，加粗\n'
    '标题字号：18-20px，加粗'
)

doc.add_heading('7.3 组件样式', level=2)
doc.add_paragraph(
    '• 容器：圆角 8px，阴影 0 1px 2px rgba(0,28,36,0.15)\n'
    '• 按钮：圆角 4px，Primary 蓝色背景白色文字，Normal 白色背景灰色边框\n'
    '• Badge：圆角 4px，不同颜色对应不同状态\n'
    '• 表格：表头灰色背景，行悬停浅灰，底部边框分割'
)

doc.add_paragraph('')

# ============================================================
# 8. 技术实现方案
# ============================================================
doc.add_heading('8. 技术实现方案', level=1)

doc.add_heading('8.1 技术栈', level=2)
doc.add_paragraph(
    '• 纯前端实现：单个 HTML 文件，内嵌 CSS + JavaScript\n'
    '• 无需任何构建工具或框架依赖\n'
    '• 浏览器直接打开即可使用\n'
    '• 兼容 Chrome、Edge、Firefox 等主流浏览器'
)

doc.add_heading('8.2 核心代码结构', level=2)
doc.add_paragraph(
    'HTML 文件内部结构：\n'
    '1. <style> 区域：所有 CSS 样式（Cloudscape 设计令牌 + 组件样式）\n'
    '2. <body> 区域：页面 HTML 结构\n'
    '3. <script> 区域：JavaScript 逻辑\n'
    '   - fetchData()：获取数据（含 CORS 代理 + 降级逻辑）\n'
    '   - parseData()：解析 JSON，过滤溢价率 > 0 的记录\n'
    '   - classifyMarket()：根据指数名称自动分类市场\n'
    '   - renderTable()：渲染表格\n'
    '   - renderStats()：更新统计卡片\n'
    '   - sortByPremium()：排序切换逻辑\n'
    '   - filterData()：筛选与搜索逻辑'
)

doc.add_heading('8.3 数据获取流程', level=2)
doc.add_paragraph(
    '1. 用户点击【查询最新数据】\n'
    '2. 通过 CORS 代理请求集思录 API\n'
    '3. 如果代理请求失败，尝试备用代理\n'
    '4. 如果所有代理都失败，使用内置示例数据并提示用户\n'
    '5. 解析返回的 JSON 数据\n'
    '6. 过滤 discount_rt > 0 的记录\n'
    '7. 按溢价率降序排列\n'
    '8. 渲染到页面'
)

doc.add_paragraph('')

# ============================================================
# 9. 文件交付物
# ============================================================
doc.add_heading('9. 文件交付物', level=1)
doc.add_paragraph(
    '最终交付一个独立的 HTML 文件：\n\n'
    '文件名：lof_premium_monitor.html\n'
    '存放路径：C:\\Users\\wuminqia\\Documents\\MQ\\\n\n'
    '该文件可直接用浏览器打开使用，无需安装任何依赖。'
)

doc.add_paragraph('')

# ============================================================
# 10. 后续可扩展方向
# ============================================================
doc.add_heading('10. 后续可扩展方向（可选）', level=1)
extensions = [
    '自动刷新：设置定时器，每 N 分钟自动刷新数据',
    '历史溢价率趋势图：记录每次查询的溢价率，绘制趋势折线图',
    '溢价率预警：设置阈值，溢价率超过阈值时高亮或弹窗提醒',
    '导出功能：支持将当前表格数据导出为 CSV / Excel',
    '折价基金展示：增加 Tab 页展示折价率较高的基金',
]
for ext in extensions:
    doc.add_paragraph(ext, style='List Bullet')

doc.add_paragraph('')

# ============================================================
# 页脚
# ============================================================
doc.add_paragraph('')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run('— 文档结束 —')
r.font.color.rgb = RGBColor(150, 150, 150)
r.font.size = Pt(10)

# 保存
doc.save(output_path)
print(f'文档已保存到: {output_path}')

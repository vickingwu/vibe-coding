# -*- coding: utf-8 -*-
# Word document generation - business value report
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
st = doc.styles['Normal']
st.font.name = '微软雅黑'; st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

t = doc.add_heading('AI Summary Report 业务价值汇报', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t.runs: r.font.color.rgb = RGBColor(0,51,102)
doc.add_paragraph('')

it = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
for i,(k,v) in enumerate([('文档类型','业务价值汇报'),('日期','2026-03-10'),('数据来源',f'VoS QS Data ({total} entries)'),('状态','For Leadership Review')]):
    it.rows[i].cells[0].text = k; it.rows[i].cells[1].text = v
doc.add_paragraph('')

doc.add_heading('1. 这份Report是什么', level=1)
doc.add_paragraph(f'基于{total}条真实VoS数据自动生成的AI Analysis Summary Report，覆盖{unique_sellers}个卖家，时间{date_min}至{date_max}。自动完成主题聚类（6个主题）、情感分析、VoS类型分布，以及基于卖家属性的多维度De-averaged Insights。')
doc.add_paragraph('')

doc.add_heading('2. 为什么需要这份Report', level=1)
for t2,d in [('数据到洞察的转化效率极低','Topic Owner手动分析一份报告需2-3天'),('聚合分析掩盖关键差异','不同GMS Band、Marketplace的卖家对同一话题反应截然不同'),('缺乏结构化分析框架','不同分析师产出不一致')]:
    p = doc.add_paragraph(); r = p.add_run(f'{t2}：'); r.bold = True; p.add_run(d)
doc.add_paragraph('')

doc.add_heading('3. De-averaged Insights核心价值', level=1)
doc.add_paragraph('每个Band的De-averaged卡片包含以下维度：')
dt = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
for i,(a,b,c) in enumerate([('分析维度','业务含义','决策价值'),('VoS Type分布','不同Band卖家反馈类型差异','了解核心诉求类型'),('情感分布','情感倾向差异','识别重点关注群体'),('Marketplace分布','站点分布差异','识别区域差异化需求'),('Domain分布','业务领域差异','精准定位核心关切'),('浓度分析','各Band在主题中占比','评估参与度'),('数据质量','高质量数据占比','评估结论可靠性')]):
    dt.rows[i].cells[0].text = a; dt.rows[i].cells[1].text = b; dt.rows[i].cells[2].text = c
doc.add_paragraph('')

doc.add_heading('4. 真实数据验证的关键发现', level=1)
findings = []
for nm,_ in theme_sorted[:4]:
    a = theme_analysis[nm]
    for band,bd in a['band_rich'].items():
        bn = bd['sentiment'].get('Negative',0)
        if bd['count']>=5 and abs(bn-a['neg_pct'])>8:
            dr = '高于' if bn>a['neg_pct'] else '低于'
            mp_s = ', '.join(f'{m}({round(c2/bd["count"]*100)}%)' for m,c2 in bd.get('mp_top',[])[:3])
            findings.append(f'"{nm}"中{band_labels.get(band,"Band "+band)}负面{bn}%，{dr}均值{a["neg_pct"]}%。主要站点：{mp_s}')
for f2 in findings[:5]: doc.add_paragraph(f2, style='List Bullet')
doc.add_paragraph('')

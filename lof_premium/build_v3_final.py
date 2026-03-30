# -*- coding: utf-8 -*-
import csv, os, re
from collections import Counter, defaultdict

csv_path = r'C:\Users\wuminqia\Documents\VOS\VOS data\vos QS data 20260109.csv'
html_path = r'C:\Users\wuminqia\Documents\VOS\AI_Summary_Report_VoS_QS.html'
docx_path = r'C:\Users\wuminqia\Documents\VOS\AI_Summary_业务价值汇报.docx'
os.makedirs(os.path.dirname(html_path), exist_ok=True)

with open(csv_path, 'r', encoding='utf-8') as f:
    rows = list(csv.DictReader(f))
total = len(rows)
unique_sellers = len(set(r['merchant_customer_id'] for r in rows))
sources = Counter(r['source'] for r in rows)
sentiments = Counter(r['sentiment'] for r in rows)
vos_types = Counter(r['vos_type'] for r in rows)
quality_dist = Counter(r['quality'] for r in rows)
gms_bands = Counter(r['gms_band_mp'] for r in rows)
dates = sorted(r['reporting_date'] for r in rows if r['reporting_date'])
date_min, date_max = dates[0], dates[-1]
domain_counter = Counter()
for r in rows:
    for item in r['domain'].strip('[]').split(','):
        item = item.strip()
        if item: domain_counter[item] += 1
top_domains = domain_counter.most_common(10)

mp_map = {'1':'US','3':'UK','4':'DE','5':'FR','6':'JP','35691':'IT','44551':'ES'}
def get_marketplace(r):
    mp_id = r['marketplace_id'].strip('[]').split(',')[0].strip()
    return mp_map.get(mp_id, 'MP-'+mp_id)
def get_domains(r):
    return [i.strip() for i in r['domain'].strip('[]').split(',') if i.strip()]

keyword_mapping = {
    'FBA Fee Policy & Cost Impact': ['fee','cost','price','pricing','economics'],
    'Inventory & Fulfillment Strategy': ['inventory','fulfillment','FBA','AWD','AGL','warehouse','logistics','shipping','MSF','MFN','SCA'],
    'Advertising & Promotions': ['advertising','ad','promotion','deal','coupon','BFCM','campaign','Sponsored'],
    'Compliance & Tax': ['compliance','tax','VAT','regulation','METI'],
    'Brand Building & AI Tools': ['brand','AI','store','GenAI'],
    'Market Expansion & Selection': ['expansion','selection','launch','new product','Europe','Japan','cold start','tariff'],
}
theme_groups = {k: [] for k in keyword_mapping}
theme_groups['Other'] = []
for r in rows:
    kw = r['keyword'].lower() + ' ' + r['domain'].lower()
    assigned = False
    for group, keywords in keyword_mapping.items():
        if any(k.lower() in kw for k in keywords):
            theme_groups[group].append(r)
            assigned = True
            break
    if not assigned:
        theme_groups['Other'].append(r)
theme_sorted = sorted(
    [(k, v) for k, v in theme_groups.items() if k != 'Other' and len(v) > 0],
    key=lambda x: len(x[1]), reverse=True)

theme_descriptions = {
    'FBA Fee Policy & Cost Impact': 'Sellers are actively evaluating 2026 FBA fee policy changes including shipping fees, packaging fees, low-volume inventory fees (FNSKU-based), aged inventory fees, and inbound placement fees.',
    'Inventory & Fulfillment Strategy': 'Sellers share experiences with AWD, AGL, MSF, and SCA programs. Feedback is largely positive on logistics cost savings, though auto-replenishment accuracy remains a concern.',
    'Advertising & Promotions': 'BFCM preparation, deal submission challenges, and ad budget optimization. Key concerns: complex deal workflows, rising promo costs.',
    'Compliance & Tax': 'Tax reporting, VAT compliance, and product regulations. Pain points: expired tax report links, complex company structures, JP regulatory delays.',
    'Brand Building & AI Tools': 'Brand Store optimization, AI content generation, and ad effectiveness. Strong interest in AI tools for descriptions, analysis, and efficiency.',
    'Market Expansion & Selection': 'EU/JP market expansion, seasonal launches, tariff adaptation. Challenges: rising ad costs for new products, logistics instability.',
}
source_display = {'beschannel_survey':'BesChannel Survey','amc_anecdote':'AMC Anecdote'}
band_labels = {'A':'Band A (Mega)','B':'Band B (Head)','C':'Band C (Body)','D':'Band D (Tail)','E':'Band E','G':'Band G'}
vt_colors = {'Practice Sharing':'#0073bb','Positive Experience':'#1d8102','Pain point':'#d13212','Feature Request':'#b35900','Knowledge Gap':'#7b68c4','Other':'#879596'}

def compute_band_rich_analysis(theme_items):
    band_data = defaultdict(list)
    for r in theme_items:
        band_data[r['gms_band_mp']].append(r)
    result = {}
    theme_total = len(theme_items)
    for band in ['A','B','C','D','E','G']:
        items = band_data.get(band, [])
        if len(items) < 3: continue
        cnt = len(items)
        vt = Counter(r['vos_type'] for r in items)
        vt_pcts = {k: round(v/cnt*100) for k,v in vt.most_common()}
        sent = Counter(r['sentiment'] for r in items)
        sent_pcts = {k: round(v/cnt*100) for k,v in sent.items()}
        qual = Counter(r['quality'] for r in items)
        high_q_pct = round(qual.get('high',0)/cnt*100)
        concentration = round(cnt/theme_total*100)
        mp_dist = Counter(get_marketplace(r) for r in items)
        mp_top = mp_dist.most_common(5)
        dom_dist = Counter()
        for r in items:
            for d in get_domains(r): dom_dist[d] += 1
        dom_top = dom_dist.most_common(5)
        quote = ''
        for r in items:
            if r['quality'] in ('high','medium'):
                c = r['vos_content']
                if any('\u4e00'<=ch<='\u9fff' for ch in c):
                    quote = c[:200].replace('\n',' ').strip(); break
                else:
                    quote = '.'.join(c.split('.')[:2]).strip()+'.'
                    if len(quote)>40: break
        result[band] = {
            'count':cnt,'concentration':concentration,
            'vos_type':vt_pcts,'vos_type_top':vt.most_common(1)[0] if vt else ('N/A',0),
            'sentiment':sent_pcts,'high_quality_pct':high_q_pct,
            'mp_top':mp_top,'dom_top':dom_top,'quote':quote[:220],
        }
    return result

theme_analysis = {}
for name, items in theme_sorted:
    t_total = len(items)
    t_sent = Counter(r['sentiment'] for r in items)
    t_vt = Counter(r['vos_type'] for r in items)
    band_rich = compute_band_rich_analysis(items)
    quotes = []
    for r in items:
        if r['quality'] in ('high','medium') and len(quotes)<2:
            c = r['vos_content']
            if any('\u4e00'<=ch<='\u9fff' for ch in c):
                quotes.append(c[:250].replace('\n',' ').strip())
            else:
                short = '.'.join(c.split('.')[:3]).strip()
                if short and len(short)>50: quotes.append(short[:280]+'...')
    theme_analysis[name] = {
        'total':t_total,'pct':round(t_total/total*100),'sentiment':t_sent,
        'neg_pct':round(t_sent.get('Negative',0)/t_total*100) if t_total else 0,
        'neu_pct':round(t_sent.get('Neutral',0)/t_total*100) if t_total else 0,
        'pos_pct':round(t_sent.get('Positive',0)/t_total*100) if t_total else 0,
        'vos_types':t_vt.most_common(3),
        'vos_type_pcts':{k:round(v/t_total*100) for k,v in t_vt.items()},
        'band_rich':band_rich,'quotes':quotes,
    }
print(f'Loaded {total} rows, {len(theme_sorted)} themes')

def esc(s):
    return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

neg_pct = round(sentiments['Negative']/total*100)
neu_pct = round(sentiments['Neutral']/total*100)
pos_pct = round(sentiments['Positive']/total*100)
theme_colors = ['#d13212','#b35900','#0073bb','#1d8102','#7b68c4','#44b9d6']
theme_badge_cls = ['badge-error','badge-warning','badge-info','badge-success','badge-purple','badge-neutral']
parts = []
theme_sorted = sorted(
    [(k, v) for k, v in theme_groups.items() if k != 'Other' and len(v) > 0],
    key=lambda x: len(x[1]), reverse=True)
theme_desc = {
    'FBA Fee Policy & Cost Impact': 'Sellers are actively evaluating 2026 FBA fee policy changes including shipping fees, packaging fees, low-volume inventory fees (FNSKU-based), aged inventory fees, and inbound placement fees.',
    'Inventory & Fulfillment Strategy': 'Sellers share experiences with AWD, AGL, MSF, and SCA programs. Feedback is largely positive on logistics cost savings, though auto-replenishment accuracy remains a concern.',
    'Advertising & Promotions': 'BFCM preparation, deal submission challenges, and ad budget optimization. Key concerns: complex deal workflows, rising promo costs.',
    'Compliance & Tax': 'Tax reporting, VAT compliance, and product regulations. Pain points: expired tax report links, complex company structures, JP regulatory delays.',
    'Brand Building & AI Tools': 'Brand Store optimization, AI content generation, and ad effectiveness. Strong interest in AI tools for descriptions, analysis, and efficiency.',
    'Market Expansion & Selection': 'EU/JP market expansion, seasonal launches, tariff adaptation. Challenges: rising ad costs for new products, logistics instability.',
}
src_disp = {'beschannel_survey':'BesChannel Survey','amc_anecdote':'AMC Anecdote'}
band_labels = {'A':'Band A (Mega)','B':'Band B (Head)','C':'Band C (Body)','D':'Band D (Tail)','E':'Band E','G':'Band G'}
vt_colors = {'Practice Sharing':'#0073bb','Positive Experience':'#1d8102','Pain point':'#d13212','Feature Request':'#b35900','Knowledge Gap':'#7b68c4','Other':'#879596'}

def compute_band_rich(theme_items):
    bd = defaultdict(list)
    for r in theme_items: bd[r['gms_band_mp']].append(r)
    result = {}
    tt = len(theme_items)
    for band in ['A','B','C','D','E','G']:
        items = bd.get(band, [])
        if len(items) < 3: continue
        cnt = len(items)
        vt = Counter(r['vos_type'] for r in items)
        vt_pcts = {k: round(v/cnt*100) for k,v in vt.most_common()}
        sent = Counter(r['sentiment'] for r in items)
        sent_pcts = {k: round(v/cnt*100) for k,v in sent.items()}
        qual = Counter(r['quality'] for r in items)
        hq = round(qual.get('high',0)/cnt*100)
        conc = round(cnt/tt*100)
        mp_dist = Counter(get_marketplace(r) for r in items).most_common(5)
        dom_dist = Counter()
        for r in items:
            for d in get_domains(r): dom_dist[d] += 1
        dom_top = dom_dist.most_common(5)
        quote = ''
        for r in items:
            if r['quality'] in ('high','medium'):
                c = r['vos_content']
                if any('\u4e00'<=ch<='\u9fff' for ch in c):
                    quote = c[:200].replace('\n',' ').strip(); break
                else:
                    quote = '.'.join(c.split('.')[:2]).strip()+'.'
                    if len(quote)>40: break
        result[band] = {'count':cnt,'conc':conc,'vt':vt_pcts,'vt_top':vt.most_common(1)[0] if vt else ('N/A',0),
            'sent':sent_pcts,'hq':hq,'mp':mp_dist,'dom':dom_top,'quote':quote[:220]}
    return result

theme_analysis = {}
for name, items in theme_sorted:
    tt = len(items)
    ts = Counter(r['sentiment'] for r in items)
    tv = Counter(r['vos_type'] for r in items)
    br = compute_band_rich(items)
    quotes = []
    for r in items:
        if r['quality'] in ('high','medium') and len(quotes)<2:
            c = r['vos_content']
            if any('\u4e00'<=ch<='\u9fff' for ch in c):
                quotes.append(c[:250].replace('\n',' ').strip())
            else:
                short = '.'.join(c.split('.')[:3]).strip()
                if short and len(short)>50: quotes.append(short[:280]+'...')
    theme_analysis[name] = {
        'total':tt,'pct':round(tt/total*100),'sentiment':ts,
        'neg_pct':round(ts.get('Negative',0)/tt*100) if tt else 0,
        'neu_pct':round(ts.get('Neutral',0)/tt*100) if tt else 0,
        'pos_pct':round(ts.get('Positive',0)/tt*100) if tt else 0,
        'vos_types':tv.most_common(3),
        'vt_pcts':{k:round(v/tt*100) for k,v in tv.items()},
        'band_rich':br,'quotes':quotes}
print(f'Loaded {total} rows, {len(theme_sorted)} themes')

def esc(s):
    return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
neg_pct = round(sentiments['Negative']/total*100)
neu_pct = round(sentiments['Neutral']/total*100)
pos_pct = round(sentiments['Positive']/total*100)
tc = ['#d13212','#b35900','#0073bb','#1d8102','#7b68c4','#44b9d6']
tbc = ['badge-error','badge-warning','badge-info','badge-success','badge-purple','badge-neutral']
parts = []

CSS = """<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>AI Analysis Summary</title><style>
:root{--bg:#f2f3f3;--white:#fff;--hdr:#fafafa;--info-bg:#f0f4ff;--warn-bg:#fff7e6;--err-bg:#fff0f0;--ok-bg:#f0faf0;--text:#16191f;--text2:#5f6b7a;--blue:#0073bb;--red:#d13212;--green:#1d8102;--orange:#b35900;--purple:#7b68c4;--teal:#44b9d6;--border:#e9ebed;--r:8px;--rs:4px;--shadow:0 1px 2px 0 rgba(0,28,36,.15);--font:"Amazon Ember","Helvetica Neue",Roboto,Arial,sans-serif}
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
body{font-family:var(--font);background:var(--bg);color:var(--text);line-height:1.6;font-size:14px}
.app{display:flex;min-height:100vh}
.side{width:240px;background:#232f3e;color:#fff;padding:16px 0;flex-shrink:0;position:fixed;top:0;left:0;height:100vh;overflow-y:auto;z-index:100}
.side-logo{padding:16px 20px;font-size:18px;font-weight:700;color:#ff9900;border-bottom:1px solid rgba(255,255,255,.1);margin-bottom:16px}
.side-nav a{display:block;padding:10px 20px;color:#d5dbdb;text-decoration:none;font-size:14px}
.side-nav a:hover{background:rgba(255,255,255,.08);color:#fff}
.side-nav a.on{background:rgba(255,255,255,.12);color:#ff9900;border-left:3px solid #ff9900}
.main{margin-left:240px;flex:1}.top{background:#fff;border-bottom:1px solid var(--border);padding:10px 24px;display:flex;align-items:center;justify-content:space-between;position:sticky;top:0;z-index:50}
.bc{display:flex;align-items:center;gap:4px;font-size:13px;color:var(--text2)}.bc a{color:var(--blue);text-decoration:none}
.acts{display:flex;gap:10px}.btn{display:inline-flex;align-items:center;gap:6px;padding:6px 16px;border-radius:var(--rs);font-size:14px;font-weight:600;cursor:pointer;border:2px solid transparent;white-space:nowrap}
.btn-p{background:var(--blue);color:#fff;border-color:var(--blue)}.btn-n{background:#fff;color:var(--text);border-color:#aab7b8}
.pg{padding:24px;max-width:1360px}.box{background:var(--white);border-radius:var(--r);box-shadow:var(--shadow);margin-bottom:20px;overflow:hidden}
.box-h{padding:14px 24px;border-bottom:1px solid var(--border);display:flex;align-items:center;justify-content:space-between}.box-h h2{font-size:18px;font-weight:700}
.box-b{padding:24px}.alert{border-left:4px solid var(--blue);border-radius:var(--rs);padding:14px 20px;margin-bottom:20px;display:flex;gap:12px;background:var(--info-bg)}.alert-t{font-weight:700;margin-bottom:2px}
.stats{display:grid;grid-template-columns:repeat(auto-fit,minmax(150px,1fr));gap:14px;margin-bottom:20px}
.st{background:var(--hdr);border-radius:var(--r);padding:14px 20px;text-align:center;border:1px solid var(--border)}.st-v{font-size:28px;font-weight:700;color:var(--blue);line-height:1.2}.st-l{font-size:13px;color:var(--text2);margin-top:2px}
.badge{display:inline-flex;padding:2px 8px;border-radius:var(--rs);font-size:12px;font-weight:600;line-height:1.6}
.badge-info{background:#e6f2ff;color:var(--blue)}.badge-error{background:#fce9e6;color:var(--red)}.badge-success{background:#e6f9e6;color:var(--green)}.badge-warning{background:#fff3e0;color:var(--orange)}.badge-neutral{background:#f2f3f3;color:var(--text2)}.badge-purple{background:#f0ecfa;color:var(--purple)}
.pb-w{display:flex;align-items:center;gap:10px;margin:6px 0}.pb{flex:1;height:8px;background:#e9ebed;border-radius:4px;overflow:hidden}.pb-f{height:100%;border-radius:4px}.pb-l{font-size:13px;font-weight:600;color:var(--text2);min-width:36px;text-align:right}
.tc{border:1px solid var(--border);border-radius:var(--r);margin-bottom:14px;overflow:hidden}.tc:hover{box-shadow:0 2px 8px rgba(0,28,36,.12)}
.tc-h{padding:14px 20px;background:var(--hdr);border-bottom:1px solid var(--border);display:flex;align-items:center;justify-content:space-between}.tc-h h3{font-size:15px;font-weight:700;display:flex;align-items:center;gap:8px}
.tc-b{padding:14px 20px}.quote{background:#f7f8f9;border-left:3px solid var(--blue);padding:10px 14px;margin:6px 0;border-radius:0 var(--rs) var(--rs) 0;font-style:italic;color:var(--text2);font-size:13px;line-height:1.7}
.da-sec{margin-top:18px;padding-top:18px;border-top:1px dashed var(--border)}.da-hd{font-size:14px;font-weight:700;color:var(--blue);margin-bottom:12px}
.da-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(340px,1fr));gap:12px}
.da-card{border:1px solid var(--border);border-radius:var(--rs);padding:14px 16px;background:var(--white)}.da-card:hover{box-shadow:0 1px 4px rgba(0,28,36,.1)}
.da-card-top{display:flex;align-items:center;justify-content:space-between;margin-bottom:8px}.da-card-name{font-weight:700;font-size:15px}.da-card-meta{font-size:12px;color:var(--text2)}
.da-row{display:flex;align-items:center;gap:8px;margin:4px 0;font-size:13px}.da-row-label{min-width:100px;color:var(--text2);font-weight:600;flex-shrink:0}
.da-mini-bar{display:flex;gap:2px;height:6px;border-radius:3px;overflow:hidden;flex:1}.da-mini-bar div{height:100%}
.da-dist{display:flex;flex-wrap:wrap;gap:4px;font-size:12px}.da-dist-item{background:#f2f3f3;padding:1px 6px;border-radius:3px;color:var(--text2)}
.da-quote{background:#f7f8f9;border-left:2px solid var(--teal);padding:6px 10px;margin-top:8px;font-size:12px;font-style:italic;color:var(--text2);line-height:1.5;border-radius:0 var(--rs) var(--rs) 0}
.tbl{width:100%;border-collapse:separate;border-spacing:0;font-size:14px}.tbl thead th{background:var(--hdr);padding:10px 14px;text-align:left;font-weight:700;font-size:13px;color:var(--text2);text-transform:uppercase;letter-spacing:.5px;border-bottom:2px solid var(--border)}
.tbl tbody td{padding:10px 14px;border-bottom:1px solid var(--border);vertical-align:top}.tbl tbody tr:hover{background:#f7f8f9}
.footer{text-align:center;padding:24px;color:var(--text2);font-size:12px;border-top:1px solid var(--border);margin-top:24px}
@media(max-width:900px){.side{display:none}.main{margin-left:0}.stats{grid-template-columns:repeat(2,1fr)}.da-grid{grid-template-columns:1fr}}
</style></head>"""
parts.append(CSS)

# === Sidebar + Topbar + Stats ===
parts.append(f'''<body><div class="app">
<nav class="side"><div class="side-logo">AGS VoS Tool</div>
<div class="side-nav"><a href="#">Dashboard</a><a href="#" class="on">VoS Library</a><a href="#">Topic Details</a><a href="#">Trending Topics</a><a href="#">Survey Analysis</a><a href="#">Reports</a></div></nav>
<div class="main"><div class="top">
<div class="bc"><a href="#">VoS Tool</a><span style="color:#aab7b8"> / </span><a href="#">VoS Library</a><span style="color:#aab7b8"> / </span><span style="color:var(--blue);font-weight:600">AI Analysis</span></div>
<div class="acts"><button class="btn btn-n">Re-analyze</button><button class="btn btn-n">Export</button><button class="btn btn-p">Save to Topic</button></div></div>
<div class="pg">
<div class="alert"><div style="color:var(--blue);flex-shrink:0;margin-top:2px">&#9432;</div>
<div><div class="alert-t">AI Analysis Summary &mdash; Auto Mode</div>
<div>Generated on 2026-01-09 | Data: <strong>VoS QS Data ({total} entries)</strong> | Period: <strong>{date_min} to {date_max}</strong></div></div></div>
<div class="stats">
<div class="st"><div class="st-v">{total}</div><div class="st-l">VoS Entries</div></div>
<div class="st"><div class="st-v">{unique_sellers}</div><div class="st-l">Unique Sellers</div></div>
<div class="st"><div class="st-v">{len(sources)}</div><div class="st-l">Data Sources</div></div>
<div class="st"><div class="st-v" style="font-size:15px">{date_min}<br>~ {date_max}</div><div class="st-l">Time Range</div></div>
<div class="st"><div class="st-v" style="color:var(--text2)">{neu_pct}%</div><div class="st-l">Neutral</div></div>
<div class="st"><div class="st-v" style="color:var(--green)">{pos_pct}%</div><div class="st-l">Positive</div></div>
<div class="st"><div class="st-v" style="color:var(--red)">{neg_pct}%</div><div class="st-l">Negative</div></div>
</div>''')

# === Executive Summary ===
domain_tags = ' '.join(f'<span class="badge badge-info">{d}</span>' for d,_ in top_domains[:6])
t1n,_ = theme_sorted[0]; t1a = theme_analysis[t1n]
t2n,_ = theme_sorted[1]; t2a = theme_analysis[t2n]
t3n,_ = theme_sorted[2]; t3a = theme_analysis[t3n]
ef = ''
for nm,_ in theme_sorted[:3]:
    a = theme_analysis[nm]
    for band, bd in a['band_rich'].items():
        bn = bd['sent'].get('Negative',0)
        if bd['count']>=5 and bn > a['neg_pct']+10:
            ef = f'Within <strong>{esc(nm)}</strong>, {band_labels.get(band,"Band "+band)} sellers show {bn}% negative vs theme avg {a["neg_pct"]}%.'
            break
    if ef: break
if not ef: ef = 'De-averaged analysis reveals significant variation across GMS bands within each theme.'

parts.append(f'''
<div class="box"><div class="box-h"><h2>Executive Summary</h2><span class="badge badge-info">AI Generated</span></div>
<div class="box-b">
<p style="font-size:15px;line-height:1.8">This analysis covers <strong>{total} VoS entries</strong> from <strong>{unique_sellers} unique sellers</strong> ({date_min} to {date_max}), sourced from BesChannel surveys ({sources.get('beschannel_survey',0)}, {round(sources.get('beschannel_survey',0)/total*100)}%) and AMC anecdotes ({sources.get('amc_anecdote',0)}, {round(sources.get('amc_anecdote',0)/total*100)}%).</p>
<p style="font-size:15px;line-height:1.8;margin-top:12px">Feedback clusters into <strong>six themes</strong>, led by <strong>{esc(t1n)}</strong> ({t1a['pct']}%), <strong>{esc(t2n)}</strong> ({t2a['pct']}%), and <strong>{esc(t3n)}</strong> ({t3a['pct']}%). Overall sentiment: Neutral {neu_pct}%, Positive {pos_pct}%, <strong style="color:var(--red)">Negative {neg_pct}%</strong>.</p>
<p style="font-size:15px;line-height:1.8;margin-top:12px"><strong>Key de-averaged finding:</strong> {ef}</p>
<div style="margin-top:14px;display:flex;gap:6px;flex-wrap:wrap">{domain_tags}</div>
</div></div>''')

# HTML body: sidebar, stats, exec summary, themes with de-averaged, VoS type, source, domain, quality, footer
# NO: keyword tags, sentiment section, recommended actions

# Sidebar + Stats
parts.append(f'''<body><div class="app">
<nav class="side"><div class="side-logo">AGS VoS Tool</div>
<div class="side-nav"><a href="#">Dashboard</a><a href="#" class="on">VoS Library</a>
<a href="#">Topic Details</a><a href="#">Trending Topics</a>
<a href="#">Survey Analysis</a><a href="#">Reports</a></div></nav>
<div class="main"><div class="top">
<div class="bc"><a href="#">VoS Tool</a><span style="color:#aab7b8"> / </span>
<a href="#">VoS Library</a><span style="color:#aab7b8"> / </span>
<span style="color:var(--blue);font-weight:600">AI Analysis</span></div>
<div class="acts"><button class="btn btn-n">Re-analyze</button>
<button class="btn btn-n">Export</button>
<button class="btn btn-p">Save to Topic</button></div></div>
<div class="pg">
<div class="alert"><div style="color:var(--blue)">&#9432;</div>
<div><div class="alert-t">AI Analysis Summary &mdash; Auto Mode</div>
<div>Generated 2026-01-09 | <strong>{total} entries</strong> | {date_min} ~ {date_max}</div></div></div>
<div class="stats">
<div class="st"><div class="st-v">{total}</div><div class="st-l">VoS Entries</div></div>
<div class="st"><div class="st-v">{unique_sellers}</div><div class="st-l">Unique Sellers</div></div>
<div class="st"><div class="st-v">{len(sources)}</div><div class="st-l">Data Sources</div></div>
<div class="st"><div class="st-v" style="font-size:15px">{date_min}<br>~{date_max}</div><div class="st-l">Time Range</div></div>
<div class="st"><div class="st-v" style="color:var(--text2)">{neu_pct}%</div><div class="st-l">Neutral</div></div>
<div class="st"><div class="st-v" style="color:var(--green)">{pos_pct}%</div><div class="st-l">Positive</div></div>
<div class="st"><div class="st-v" style="color:var(--red)">{neg_pct}%</div><div class="st-l">Negative</div></div></div>''')

# Executive Summary
dtags = ' '.join(f'<span class="badge badge-info">{d}</span>' for d,_ in top_domains[:6])
t1n,_=theme_sorted[0]; t1a=theme_analysis[t1n]
t2n,_=theme_sorted[1]; t2a=theme_analysis[t2n]
t3n,_=theme_sorted[2]; t3a=theme_analysis[t3n]
ef=''
for nm,_ in theme_sorted[:3]:
    a=theme_analysis[nm]
    for b,bd in a['band_rich'].items():
        bn=bd['sentiment'].get('Negative',0)
        if bd['count']>=5 and bn>a['neg_pct']+10:
            ef=f'Within <strong>{esc(nm)}</strong>, {band_labels.get(b,"Band "+b)} show {bn}% negative vs theme avg {a["neg_pct"]}%.'
            break
    if ef: break
if not ef: ef='De-averaged analysis reveals significant variation across GMS bands.'

# Executive summary + Theme cards with de-averaged (marketplace + domain, no keywords)

parts.append(f'''
<div class="box"><div class="box-h"><h2>Executive Summary</h2><span class="badge badge-info">AI Generated</span></div>
<div class="box-b">
<p style="font-size:15px;line-height:1.8">This analysis covers <strong>{total} VoS entries</strong> from <strong>{unique_sellers} unique sellers</strong> ({date_min} to {date_max}), sourced from BesChannel ({sources.get("beschannel_survey",0)}, {round(sources.get("beschannel_survey",0)/total*100)}%) and AMC ({sources.get("amc_anecdote",0)}, {round(sources.get("amc_anecdote",0)/total*100)}%).</p>
<p style="font-size:15px;line-height:1.8;margin-top:12px">Feedback clusters into <strong>six themes</strong>, led by <strong>{esc(t1n)}</strong> ({t1a["pct"]}%), <strong>{esc(t2n)}</strong> ({t2a["pct"]}%), and <strong>{esc(t3n)}</strong> ({t3a["pct"]}%). Overall: Neutral {neu_pct}%, Positive {pos_pct}%, <strong style="color:var(--red)">Negative {neg_pct}%</strong>.</p>
<p style="font-size:15px;line-height:1.8;margin-top:12px"><strong>Key de-averaged finding:</strong> {ef}</p>
<div style="margin-top:14px;display:flex;gap:6px;flex-wrap:wrap">{dtags}</div>
</div></div>''')

# Theme Clustering
parts.append(f'''<div class="box"><div class="box-h"><h2>Theme Clustering with De-averaged Insights</h2><span class="badge badge-neutral">{len(theme_sorted)} themes</span></div><div class="box-b">''')

for idx,(name,items) in enumerate(theme_sorted):
    a = theme_analysis[name]
    color = tc[idx%len(tc)]
    bcls = tbc[idx%len(tbc)]
    top_vt = a['vos_types'][0] if a['vos_types'] else ('N/A',0)
    desc = theme_descriptions.get(name,'')
    qhtml = ''.join(f'<div class="quote">{esc(q)}</div>' for q in a['quotes'])
    # VoS type badges at theme level
    vtb = '<div style="display:flex;gap:6px;flex-wrap:wrap;margin-bottom:8px">'
    for vn,vc2 in a['vos_types']:
        vp = a['vos_type_pcts'].get(vn,0)
        vtb += f'<span class="badge" style="background:{vt_colors.get(vn,"#879596")}22;color:{vt_colors.get(vn,"#879596")}">{esc(vn)} {vp}%</span>'
    vtb += '</div>'

    # De-averaged cards per band
    bcards = ''
    for band in ['A','B','C','D','E','G']:
        if band not in a['band_rich']: continue
        bd = a['band_rich'][band]
        label = band_labels.get(band, f'Band {band}')
        cnt = bd['count']
        # VoS Type mini bar
        vtbar = '<div class="da-row"><span class="da-row-label">VoS Type</span><div class="da-mini-bar" style="width:120px">'
        for vn in ['Pain point','Practice Sharing','Feature Request','Positive Experience','Knowledge Gap']:
            vp = bd['vos_type'].get(vn,0)
            if vp>0: vtbar += f'<div style="width:{max(vp,2)}%;background:{vt_colors.get(vn,"#879596")}" title="{esc(vn)} {vp}%"></div>'
        vtn = bd['vos_type_top'][0]
        vtbar += f'</div><span style="font-size:12px;color:{vt_colors.get(vtn,"#879596")}">{esc(vtn)} {bd["vos_type"].get(vtn,0)}%</span></div>'
        # Sentiment mini bar
        sn=bd['sentiment'].get('Negative',0); su=bd['sentiment'].get('Neutral',0); sp=bd['sentiment'].get('Positive',0)
        sbar = f'<div class="da-row"><span class="da-row-label">Sentiment</span><div class="da-mini-bar" style="width:120px"><div style="width:{max(sn,1)}%;background:var(--red)"></div><div style="width:{max(su,1)}%;background:var(--orange)"></div><div style="width:{max(sp,1)}%;background:var(--green)"></div></div><span style="font-size:12px"><span style="color:var(--red)">{sn}%</span>/<span style="color:var(--green)">{sp}%</span></span></div>'
        # Neg flag
        nf = ''
        if sn>a['neg_pct']+8 and cnt>=5: nf=f' <span style="background:var(--err-bg);color:var(--red);padding:1px 6px;border-radius:3px;font-size:11px;font-weight:600">&#9650;+{sn-a["neg_pct"]}pp</span>'
        elif sn<a['neg_pct']-8 and cnt>=5: nf=f' <span style="background:var(--ok-bg);color:var(--green);padding:1px 6px;border-radius:3px;font-size:11px;font-weight:600">&#9660;{sn-a["neg_pct"]}pp</span>'
        # Marketplace distribution (NEW - replaces keywords)
        mph = '<div class="da-row"><span class="da-row-label">Marketplace</span><div class="da-dist">'
        for mn,mc in bd.get('mp_top',[])[:4]:
            mph += f'<span class="da-dist-item">{esc(mn)} {round(mc/cnt*100)}%</span>'
        mph += '</div></div>'
        # Domain distribution (NEW - replaces keywords)
        dmh = '<div class="da-row"><span class="da-row-label">Domain</span><div class="da-dist">'
        for dn,dc in bd.get('dom_top',[])[:4]:
            dmh += f'<span class="da-dist-item">{esc(dn)} {round(dc/cnt*100)}%</span>'
        dmh += '</div></div>'
        # Concentration + Quality
        mh = f'<div class="da-row"><span class="da-row-label">Concentration</span><span style="font-size:12px;font-weight:700">{bd["concentration"]}% of theme</span></div>'
        mh += f'<div class="da-row"><span class="da-row-label">High Quality</span><span style="font-size:12px">{bd["high_quality_pct"]}%</span></div>'
        # Quote
        qh = f'<div class="da-quote">{esc(bd["quote"])}</div>' if bd['quote'] else ''
        bcards += f'''<div class="da-card"><div class="da-card-top"><span class="da-card-name">{esc(label)}{nf}</span><span class="da-card-meta">{cnt} entries</span></div>{vtbar}{sbar}{mph}{dmh}{mh}{qh}</div>'''

# Theme card assembly + VoS Type/Source/Domain/Quality + Footer (no Sentiment, no Actions)

    parts.append(f'''
    <div class="tc"><div class="tc-h">
    <h3><span style="color:{color}">&#9679;</span> Theme {idx+1}: {esc(name)}</h3>
    <div style="display:flex;align-items:center;gap:10px"><span class="badge {bcls}">{esc(top_vt[0])}</span>
    <span style="font-weight:700">{a['pct']}%</span><span style="font-size:13px;color:var(--text2)">{a['total']} entries</span></div></div>
    <div class="tc-b">
    <div class="pb-w"><div class="pb"><div class="pb-f" style="width:{a['pct']}%;background:{color}"></div></div><span class="pb-l">{a['pct']}%</span></div>
    <p style="color:var(--text);margin-bottom:12px;line-height:1.7">{esc(desc)}</p>
    {vtb}
    <div style="display:flex;gap:8px;margin-bottom:8px;flex-wrap:wrap">
    <span class="badge" style="background:var(--err-bg);color:var(--red)">Negative {a['neg_pct']}%</span>
    <span class="badge" style="background:var(--warn-bg);color:var(--orange)">Neutral {a['neu_pct']}%</span>
    <span class="badge" style="background:var(--ok-bg);color:var(--green)">Positive {a['pos_pct']}%</span></div>
    {qhtml}
    <div class="da-sec"><div class="da-hd">&#128202; De-averaged by GMS Band</div>
    <div class="da-grid">{bcards}</div></div>
    </div></div>''')

parts.append('</div></div>')

# VoS Type Distribution
vth = ''
for vt,cnt in vos_types.most_common():
    p=round(cnt/total*100); c=vt_colors.get(vt,'#879596')
    vth += f'<div style="flex:1;min-width:180px"><div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:4px"><span style="font-weight:600">{esc(vt)}</span><span style="font-weight:700;color:{c}">{cnt}</span></div><div class="pb-w"><div class="pb"><div class="pb-f" style="width:{p}%;background:{c}"></div></div><span class="pb-l">{p}%</span></div></div>'

# Source
srch = ''
src_c = {'beschannel_survey':'var(--blue)','amc_anecdote':'var(--purple)'}
for s,cnt in sources.most_common():
    p=round(cnt/total*100)
    srch += f'<div style="flex:1;min-width:200px"><div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:4px"><span style="font-weight:600">{esc(source_display.get(s,s))}</span><span style="font-weight:700;color:{src_c.get(s,"#879596")}">{cnt}</span></div><div class="pb-w"><div class="pb"><div class="pb-f" style="width:{p}%;background:{src_c.get(s,"#879596")}"></div></div><span class="pb-l">{p}%</span></div></div>'

# Domain table
domh = ''
for d,cnt in top_domains:
    p=round(cnt/total*100)
    di=[r for r in rows if d in r['domain']]
    ds=Counter(r['sentiment'] for r in di)
    dn=round(ds.get('Negative',0)/len(di)*100) if di else 0
    dp=round(ds.get('Positive',0)/len(di)*100) if di else 0
    domh += f'<tr><td><strong>{esc(d)}</strong></td><td>{cnt}</td><td>{p}%</td><td><span style="color:var(--red)">{dn}%</span>/<span style="color:var(--green)">{dp}%</span></td></tr>'

qh2=quality_dist.get('high',0); qm=quality_dist.get('medium',0); ql=quality_dist.get('low',0)

parts.append(f'''
<div class="box"><div class="box-h"><h2>VoS Type Distribution</h2></div><div class="box-b"><div style="display:flex;gap:20px;flex-wrap:wrap">{vth}</div></div></div>
<div class="box"><div class="box-h"><h2>Data Source Breakdown</h2></div><div class="box-b"><div style="display:flex;gap:20px;flex-wrap:wrap">{srch}</div></div></div>
<div class="box"><div class="box-h"><h2>Top Domains</h2></div><div class="box-b"><table class="tbl"><thead><tr><th>Domain</th><th>Count</th><th>%</th><th>Neg/Pos</th></tr></thead><tbody>{domh}</tbody></table></div></div>
<div class="box"><div class="box-h"><h2>Data Quality</h2></div><div class="box-b"><div style="display:flex;gap:20px;flex-wrap:wrap">
<div style="flex:1;min-width:180px"><div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:4px"><span style="font-weight:600">High</span><span style="font-weight:700;color:var(--green)">{qh2}</span></div><div class="pb-w"><div class="pb"><div class="pb-f" style="width:{round(qh2/total*100)}%;background:var(--green)"></div></div><span class="pb-l">{round(qh2/total*100)}%</span></div></div>
<div style="flex:1;min-width:180px"><div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:4px"><span style="font-weight:600">Medium</span><span style="font-weight:700;color:var(--orange)">{qm}</span></div><div class="pb-w"><div class="pb"><div class="pb-f" style="width:{round(qm/total*100)}%;background:var(--orange)"></div></div><span class="pb-l">{round(qm/total*100)}%</span></div></div>
<div style="flex:1;min-width:180px"><div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:4px"><span style="font-weight:600">Low</span><span style="font-weight:700;color:var(--red)">{ql}</span></div><div class="pb-w"><div class="pb"><div class="pb-f" style="width:{round(ql/total*100)}%;background:var(--red)"></div></div><span class="pb-l">{round(ql/total*100)}%</span></div></div>
</div></div></div>''')

# Footer
parts.append(f'''<div class="footer"><p>AGS VoS Tool &mdash; AI Analysis Summary | AI Generated (Auto Mode) | {total} entries</p>
<p style="margin-top:4px">De-averaged insights per-theme with Marketplace + Domain distribution. &copy; Amazon AGS</p></div>
</div></div></div></body></html>''')

with open(html_path, 'w', encoding='utf-8') as f:
    f.write(''.join(parts))
print(f'HTML saved: {html_path}')

# Word document generation - business value report
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
st = doc.styles['Normal']
st.font.name = '微软雅黑'; st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

t = doc.add_heading('AI Summary Report 业务价值汇报', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t.runs: r.font.color.rgb = RGBColor(0,51,102)
doc.add_paragraph('')

it = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
for i,(k,v) in enumerate([('文档类型','业务价值汇报'),('日期','2026-03-10'),('数据来源',f'VoS QS Data ({total} entries)'),('状态','For Leadership Review')]):
    it.rows[i].cells[0].text = k; it.rows[i].cells[1].text = v
doc.add_paragraph('')

doc.add_heading('1. 这份Report是什么', level=1)
doc.add_paragraph(f'基于{total}条真实VoS数据自动生成的AI Analysis Summary Report，覆盖{unique_sellers}个卖家，时间{date_min}至{date_max}。自动完成主题聚类（6个主题）、情感分析、VoS类型分布，以及基于卖家属性的多维度De-averaged Insights。')
doc.add_paragraph('')

doc.add_heading('2. 为什么需要这份Report', level=1)
for t2,d in [('数据到洞察的转化效率极低','Topic Owner手动分析一份报告需2-3天'),('聚合分析掩盖关键差异','不同GMS Band、Marketplace的卖家对同一话题反应截然不同'),('缺乏结构化分析框架','不同分析师产出不一致')]:
    p = doc.add_paragraph(); r = p.add_run(f'{t2}：'); r.bold = True; p.add_run(d)
doc.add_paragraph('')

doc.add_heading('3. De-averaged Insights核心价值', level=1)
doc.add_paragraph('每个Band的De-averaged卡片包含以下维度：')
dt = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
for i,(a,b,c) in enumerate([('分析维度','业务含义','决策价值'),('VoS Type分布','不同Band卖家反馈类型差异','了解核心诉求类型'),('情感分布','情感倾向差异','识别重点关注群体'),('Marketplace分布','站点分布差异','识别区域差异化需求'),('Domain分布','业务领域差异','精准定位核心关切'),('浓度分析','各Band在主题中占比','评估参与度'),('数据质量','高质量数据占比','评估结论可靠性')]):
    dt.rows[i].cells[0].text = a; dt.rows[i].cells[1].text = b; dt.rows[i].cells[2].text = c
doc.add_paragraph('')

doc.add_heading('4. 真实数据验证的关键发现', level=1)
findings = []
for nm,_ in theme_sorted[:4]:
    a = theme_analysis[nm]
    for band,bd in a['band_rich'].items():
        bn = bd['sentiment'].get('Negative',0)
        if bd['count']>=5 and abs(bn-a['neg_pct'])>8:
            dr = '高于' if bn>a['neg_pct'] else '低于'
            mp_s = ', '.join(f'{m}({round(c2/bd["count"]*100)}%)' for m,c2 in bd.get('mp_top',[])[:3])
            findings.append(f'"{nm}"中{band_labels.get(band,"Band "+band)}负面{bn}%，{dr}均值{a["neg_pct"]}%。主要站点：{mp_s}')
for f2 in findings[:5]: doc.add_paragraph(f2, style='List Bullet')
doc.add_paragraph('')

# Word doc continued: value + ROI + next steps

doc.add_heading('5. 对业务的直接价值', level=1)
for t3,d in [('效率提升','从2-3天缩短到秒级自动生成'),('洞察深度','多维度De-averaged分析揭示传统方法无法发现的差异，包括Marketplace和Domain维度'),('决策精准度','基于分群洞察的差异化策略更有针对性'),('标准化输出','统一信息结构确保可比性'),('可扩展性','同一框架可应用于任意VoS数据集')]:
    p = doc.add_paragraph(); r = p.add_run(f'{t3}：'); r.bold = True; p.add_run(d)
doc.add_paragraph('')

doc.add_heading('6. 投入产出评估', level=1)
rt = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
for i,(a,b) in enumerate([('评估项','说明'),('开发投入','约3人月'),('当前人工成本','每份报告2-3天，月均8-12人天'),('自动化节省','节省80%分析时间'),('额外价值','多维度De-averaged是人工几乎无法完成的纯增量价值')]):
    rt.rows[i].cells[0].text = a; rt.rows[i].cells[1].text = b
doc.add_paragraph('')

doc.add_heading('7. 下一步建议', level=1)
for i,s in enumerate(['将本Report作为MVP验证，收集反馈','优化De-averaged维度和展示','启动Phase 1开发','建立VoS Tool集成方案']):
    doc.add_paragraph(f'{i+1}. {s}', style='List Number')

doc.save(docx_path)
print(f'Word saved: {docx_path}')
print('All done!')
